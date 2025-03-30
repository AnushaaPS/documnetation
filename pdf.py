import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
import pypandoc

# Ensure Pandoc is installed
pypandoc.download_pandoc()

# Function to format student names dynamically
def format_students(students):
    students = [f"{name.strip()} {reg.strip()}" for name, reg in students if name.strip() and reg.strip()]
    if len(students) == 1:
        return students[0]
    elif len(students) == 2:
        return f"{students[0]} & {students[1]}"
    elif len(students) > 2:
        return f"{', '.join(students[:-1])} & {students[-1]}"
    return "Unknown"

# Function to set line spacing to 1.5
def set_line_spacing(paragraph):
    p = paragraph._element
    spacing = parse_xml(r'<w:spacing w:line="360" w:lineRule="auto" %s />' % nsdecls('w'))
    p.get_or_add_pPr().append(spacing)

# Function to fill the project report template
def fill_project_report(details, template):
    doc = Document(template)  # Load template

    font_sizes = {
        "<PROJECT_NAME>": 18,
        "<STUDENT_DETAILS>": 14,
        "<STUDENT_1>": 16,
        "<REG_NO_1>": 16,
        "<STUDENT_2>": 16,
        "<REG_NO_2>": 16,
        "<STUDENT_3>": 16,
        "<REG_NO_3>": 16,
        "<STUDENT_4>": 16,
        "<REG_NO_4>": 16,
        "<DEGREE>": 16,
        "<DEPARTMENT>": 14,
        "<HOD_NAME>": 14,
        "<SUPERVISOR_NAME>": 14,
        "<DESIGNATION>": 14,
        "<DEPARTMENT_1>": 14,
        "<INDUSTRY_PERSON_NAME>": 14,
        "<INDUSTRY_PERSON_POSITION>": 14,
        "<INDUSTRY_PERSON_PRONOUN>": 14,
    }

    # Replace placeholders and apply styles
    for i, para in enumerate(doc.paragraphs):
        for key, value in details.items():
            if key in para.text:
                para.text = para.text.replace(key, value.strip())
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_sizes.get(key, 14))
        set_line_spacing(para)

        # Remove empty paragraphs at the beginning
        if i < 10 and para.text.strip() == "":
            p = para._element
            p.getparent().remove(p)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in details.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(font_sizes.get(key, 14))

    # Save document
    output = BytesIO()
    doc.save(output)
    return output

# Function to convert DOCX to PDF using pypandoc
def convert_docx_to_pdf(docx_bytes):
    temp_docx = "temp.docx"
    with open(temp_docx, "wb") as f:
        f.write(docx_bytes.getvalue())

    pdf_output = "temp.pdf"
    
    # Convert using wkhtmltopdf (no pdflatex needed)
    pypandoc.convert_file(temp_docx, "pdf", outputfile=pdf_output, extra_args=['--pdf-engine=wkhtmltopdf'])

    with open(pdf_output, "rb") as f:
        return f.read()

# Streamlit UI
st.title("Project Report Generator")
project_type = st.radio("Select Project Type", ["Internal Project", "External Project"])

with st.form("project_form"):
    project_name = st.text_input("Project Name", "")
    student_1 = st.text_input("Student 1 Name", "")
    reg_no_1 = st.text_input("Register Number 1", "")
    student_2 = st.text_input("Student 2 Name (Optional)", "")
    reg_no_2 = st.text_input("Register Number 2 (Optional)", "")
    student_3 = st.text_input("Student 3 Name (Optional)", "")
    reg_no_3 = st.text_input("Register Number 3 (Optional)", "")
    student_4 = st.text_input("Student 4 Name (Optional)", "")
    reg_no_4 = st.text_input("Register Number 4 (Optional)", "")
    degree = st.selectbox("Degree", ["BACHELOR OF ENGINEERING", "BACHELOR OF TECHNOLOGY"])
    department = st.selectbox("Department", ["COMPUTER SCIENCE AND ENGINEERING", "INFORMATION TECHNOLOGY", "ELECTRONICS AND COMMUNICATION ENGINEERING"])
    hod_name = st.text_input("HoD Name", "")
    hod_gender = st.radio("HoD Gender", ["Male", "Female"])
    supervisor_name = st.text_input("Supervisor Name", "")
    supervisor_gender = st.radio("Supervisor Gender", ["Male", "Female"])
    supervisor_designation = st.selectbox("Supervisor Designation", ["Assistant Professor", "Associate Professor", "Professor"])
    department_hod_supervisor = st.selectbox("Department of HoD & Supervisor", ["Computer Science and Engineering", "Artificial Intelligence And Data Science", "Information Technology", "Electronics And Communication Engineering","Electrical And Electronics Engineering","Bio-Technology", "Mechanical Engineering","Mechatronics Engineering", "Civil Engineering"])

    if project_type == "External Project":
        industry_name = st.text_input("Industry Name", "")
        industry_person_name = st.text_input("Industry Person Name", "")
        industry_person_position = st.text_input("Industry Person Position", "")
        industry_person_gender = st.radio("Industry Person Gender", ["Male", "Female"])

    submitted = st.form_submit_button("Generate Report")

if submitted:
    students_list = [(student_1, reg_no_1), (student_2, reg_no_2), (student_3, reg_no_3), (student_4, reg_no_4)]
    formatted_students = format_students(students_list)

    details = {
        "<PROJECT_NAME>": project_name, "<STUDENT_DETAILS>": formatted_students,
        "<STUDENT_1>": student_1, "<REG_NO_1>": reg_no_1, "<STUDENT_2>": student_2,
        "<REG_NO_2>": reg_no_2, "<STUDENT_3>": student_3, "<REG_NO_3>": reg_no_3,
        "<STUDENT_4>": student_4, "<REG_NO_4>": reg_no_4, "<DEGREE>": degree,
        "<DEPARTMENT>": department, "<HOD_NAME>": hod_name, "<SUPERVISOR_NAME>": supervisor_name,
        "<DESIGNATION>": supervisor_designation, "<DEPARTMENT_1>": department_hod_supervisor,
        "<HOD_PRONOUN>": "his" if hod_gender == "Male" else "her",
        "<SUPERVISOR_PRONOUN>": "his" if supervisor_gender == "Male" else "her"
    }

    if project_type == "External Project":
        details.update({
            "<INDUSTRY_NAME>": industry_name,
            "<INDUSTRY_PERSON_NAME>": industry_person_name,
            "<INDUSTRY_PERSON_POSITION>": industry_person_position,
            "<INDUSTRY_PERSON_PRONOUN>": "his" if industry_person_gender == "Male" else "her"
        })

    template = "UG Internal Project.docx" if project_type == "Internal Project" else "UG External Project.docx"
    word_output = fill_project_report(details, template)
    pdf_data = convert_docx_to_pdf(word_output)

    # Download buttons
    st.download_button("Download Report (DOCX)", word_output.getvalue(), "Project_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("Download Report (PDF)", pdf_data, "Project_Report.pdf", "application/pdf")
